from io import BytesIO
from zipfile import ZipFile

import fitz
import pytest
from docx import Document as DocxDocument

from app.services.document_processing import ExtractedBlock, chunk_blocks, extract_blocks


def test_txt_extraction_and_chunking_preserve_section():
    blocks = extract_blocks(
        "# Введение\nПервый абзац.\nВторой абзац.".encode(), "text/plain"
    )
    chunks = chunk_blocks(blocks, max_chars=40, overlap_chars=5)

    assert blocks[0].section == "Введение"
    assert chunks
    assert all(chunk["section"] == "Введение" for chunk in chunks)
    assert [chunk["chunk_index"] for chunk in chunks] == list(range(len(chunks)))


def test_docx_extraction_preserves_heading():
    stream = BytesIO()
    document = DocxDocument()
    document.add_heading("Раздел", level=1)
    document.add_paragraph("Содержимое раздела")
    document.save(stream)

    blocks = extract_blocks(
        stream.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert blocks[-1].text == "Содержимое раздела"
    assert blocks[-1].section == "Раздел"


def test_docx_extraction_preserves_table_rows_and_section():
    stream = BytesIO()
    document = DocxDocument()
    document.add_heading("Таблица", level=1)
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Понятие"
    table.cell(0, 1).text = "Описание"
    document.save(stream)

    blocks = extract_blocks(
        stream.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert blocks[-1].text == "Понятие | Описание"
    assert blocks[-1].section == "Таблица"


def test_pdf_extraction_preserves_page_number():
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "PDF content")
    content = document.tobytes()
    document.close()

    blocks = extract_blocks(content, "application/pdf")

    assert blocks[0].page == 1
    assert "PDF content" in blocks[0].text


def test_txt_extraction_builds_heading_hierarchy_and_paragraphs():
    blocks = extract_blocks(
        (
            "# Main\n"
            "First line\n"
            "continues here\n\n"
            "## Child\n"
            "Child content\n\n"
            "Sibling\n"
            "-------\n"
            "Sibling content"
        ).encode(),
        "text/plain",
    )

    assert [(item.text, item.block_type, item.heading_level) for item in blocks] == [
        ("Main", "heading", 1),
        ("First line continues here", "paragraph", None),
        ("Child", "heading", 2),
        ("Child content", "paragraph", None),
        ("Sibling", "heading", 2),
        ("Sibling content", "paragraph", None),
    ]
    assert blocks[3].section == "Main > Child"
    assert blocks[-1].section == "Main > Sibling"


def test_txt_does_not_treat_markdown_inside_fence_as_structure():
    blocks = extract_blocks(
        "# Outside\n\n```text\n# Not a heading\n```".encode(), "text/plain"
    )

    assert [item.text for item in blocks] == [
        "Outside",
        "```text # Not a heading ```",
    ]
    assert blocks[-1].block_type == "paragraph"
    assert blocks[-1].section == "Outside"


def test_docx_extraction_preserves_nested_headings_and_merged_table_cells():
    stream = BytesIO()
    document = DocxDocument()
    document.add_heading("Parent", level=1)
    document.add_heading("Child", level=2)
    document.add_paragraph("Nested content")
    table = document.add_table(rows=1, cols=2)
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.text = "Merged value"
    document.save(stream)

    blocks = extract_blocks(
        stream.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert blocks[1].section == "Parent > Child"
    assert blocks[1].heading_level == 2
    assert blocks[2].section == "Parent > Child"
    assert blocks[-1].text == "Merged value"
    assert blocks[-1].block_type == "table_row"


def test_pdf_extraction_removes_repeated_margins_and_infers_section():
    document = fitz.open()
    for page_number in range(1, 3):
        page = document.new_page()
        page.insert_text((72, 25), "Repeated header", fontsize=8)
        if page_number == 1:
            page.insert_text((72, 130), "Safety section", fontsize=18)
        page.insert_text((72, 170), f"Body content {page_number}", fontsize=11)
        page.insert_text((300, 810), str(page_number), fontsize=8)
    content = document.tobytes()
    document.close()

    blocks = extract_blocks(content, "application/pdf")

    assert "Repeated header" not in {item.text for item in blocks}
    assert "1" not in {item.text for item in blocks}
    assert "2" not in {item.text for item in blocks}
    heading = next(item for item in blocks if item.text == "Safety section")
    assert heading.block_type == "heading"
    assert heading.heading_level == 1
    assert all(
        item.section == "Safety section"
        for item in blocks
        if item.text.startswith("Body content")
    )


def test_chunking_is_deterministic_bounded_and_does_not_cross_metadata():
    alpha_words = [f"alpha{index:02d}" for index in range(24)]
    beta_words = [f"beta{index:02d}" for index in range(12)]
    blocks = [
        ExtractedBlock(text=" ".join(alpha_words), page=1, section="Alpha"),
        ExtractedBlock(text=" ".join(beta_words), page=1, section="Beta"),
    ]

    first = chunk_blocks(blocks, max_chars=64, overlap_chars=12)
    second = chunk_blocks(blocks, max_chars=64, overlap_chars=12)

    assert first == second
    assert all(0 < len(item["text"]) <= 64 for item in first)
    assert [item["chunk_index"] for item in first] == list(range(len(first)))
    assert {item["section"] for item in first} == {"Alpha", "Beta"}
    assert all(
        not ("alpha" in item["text"] and "beta" in item["text"])
        for item in first
    )
    assert all(
        item["metadata_json"]
        == {"page": item["page"], "section": item["section"]}
        for item in first
    )
    allowed_words = set(alpha_words + beta_words)
    assert all(word in allowed_words for item in first for word in item["text"].split())


def test_chunking_handles_a_long_unbroken_token_without_stalling():
    chunks = chunk_blocks(
        [ExtractedBlock(text="x" * 205)], max_chars=50, overlap_chars=10
    )

    assert len(chunks) == 5
    assert all(0 < len(item["text"]) <= 50 for item in chunks)


@pytest.mark.parametrize(
    ("max_chars", "overlap_chars"),
    [(0, 0), (-1, 0), (10, -1), (10, 10), (10, 11)],
)
def test_chunking_rejects_invalid_limits(max_chars, overlap_chars):
    with pytest.raises(ValueError):
        chunk_blocks(
            [ExtractedBlock(text="content")],
            max_chars=max_chars,
            overlap_chars=overlap_chars,
        )


def test_corrupted_supported_documents_raise_value_error():
    with pytest.raises(ValueError):
        extract_blocks(b"%PDF-1.7\nnot a complete pdf", "application/pdf")

    stream = BytesIO()
    with ZipFile(stream, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        archive.writestr("word/document.xml", "<broken")
    with pytest.raises(ValueError):
        extract_blocks(
            stream.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    with pytest.raises(ValueError):
        extract_blocks(b"\xff\xfe\xfa", "text/plain")
